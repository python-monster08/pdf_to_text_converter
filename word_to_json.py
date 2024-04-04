'''
    This Script is able to find all Koat questions from my formatted document (word file)
'''


# from docx import Document
# import json
# import re
#
# # Load the Word document
# doc = Document('sample_questions.docx')
#
# # Initialize an empty list to store all questions
# questions_list = []
#
# # Define the regex patterns
# question_pattern = re.compile(r'^(\d+)\. (.+?)\[(\d{4})\]')
# list_title_pattern = re.compile(r'List-I\s+\((.*?)\)\s+List-II\s+\((.*?)\)')
# list_I_option_pattern = re.compile(r'([ABCD])\.\s+(.+?)(?=\s+[ABCD]\.\s+|\s+\d+\.\s+|\Z)', re.DOTALL)
# list_II_option_pattern = re.compile(r'(\d+)\.\s+(.+?)(?=\s+\d+\.\s+|\Z)', re.DOTALL)
# codes_pattern = re.compile(r'Codes:\s+([ABCD\s]+)')
# options_pattern = re.compile(r'\(([abcd])\)\s+(.*?)\s*(?=\([abcd]\)|$)', re.DOTALL)
# answer_pattern = re.compile(r'Answer\s+-\s+\((a|b|c|d)\)')
# exam_name_pattern = re.compile(r'(U\.P\s+\.P\s+\.C\.S\.\s+\(.+?\)\s+\d{4})')
#
# # Loop through each paragraph in the document
# for para in doc.paragraphs:
#     text = para.text.strip()
#
#     # Skip any empty paragraphs
#     if not text:
#         continue
#
#     # Match the start of a question
#     if question_match := question_pattern.match(text):
#         question_number, initial_text, exam_year = question_match.groups()
#
#         question_dict = {
#             "question_number": question_number,
#             "question_text": initial_text.strip(),
#             "exam_year": exam_year,
#             "list_I": {"title": "", "options": {}},
#             "list_II": {"title": "", "options": {}},
#             "codes": [],
#             "options": {},
#             "answer": "",
#             "exam_name": ""
#         }
#
#         # Extract the list titles and options
#         if list_match := list_title_pattern.search(text):
#             list_I_title, list_II_title = list_match.groups()
#             question_dict["list_I"]["title"] = list_I_title.strip()
#             question_dict["list_II"]["title"] = list_II_title.strip()
#
#             # Extract options for List I
#             for match in list_I_option_pattern.finditer(text):
#                 letter, option = match.groups()
#                 question_dict["list_I"]["options"][letter] = option.strip()
#
#             # Extract options for List II, removing extra content
#             for match in list_II_option_pattern.finditer(text[list_match.end():]):
#                 number, option = match.groups()
#                 option = re.sub(r'\n.*', '', option)  # Remove any content after a newline
#                 question_dict["list_II"]["options"][number] = option.strip()
#
#         # Extract codes
#         if codes_match := codes_pattern.search(text):
#             question_dict["codes"] = codes_match.group(1).split()
#
#         # Extract options with their codes, ensuring no extraneous text is included
#         for match in options_pattern.finditer(text):
#             code, option = match.groups()
#             option = re.sub(r'\n.*', '', option)  # Remove any content after a newline
#             question_dict["options"][code] = option.strip()
#
#         # Extract the correct answer
#         if answer_match := answer_pattern.search(text):
#             question_dict["answer"] = answer_match.group(1)
#
#         # Extract the exam name
#         if exam_name_match := exam_name_pattern.search(text):
#             question_dict["exam_name"] = exam_name_match.group(1)
#
#         questions_list.append(question_dict)
#
# # Convert the list of questions to JSON format
# questions_json = json.dumps(questions_list, indent=4, ensure_ascii=False)
# print(questions_json)

#
# import json
# import re
# from docx import Document
# from openpyxl import Workbook
# def parse_word_document(doc_path):
#     doc = Document(doc_path)
#     questions_list = []
#
#     question_pattern = re.compile(r'^(\d+)\. (.+?)\[(\d{4})\]')
#     list_title_pattern = re.compile(r'List-I\s+\((.*?)\)\s+List-II\s+\((.*?)\)')
#     list_I_option_pattern = re.compile(r'([ABCD])\.\s+(.+?)(?=\s+[ABCD]\.\s+|\s+\d+\.\s+|\Z)', re.DOTALL)
#     list_II_option_pattern = re.compile(r'(\d+)\.\s+(.+?)(?=\s+\d+\.\s+|\Z)', re.DOTALL)
#     codes_pattern = re.compile(r'Codes:\s+([ABCD\s]+)')
#     options_pattern = re.compile(r'\(([abcd])\)\s+(.*?)\s*(?=\([abcd]\)|$)', re.DOTALL)
#     answer_pattern = re.compile(r'Answer\s+-\s+\((a|b|c|d)\)')
#     # Adjusted regex to separately capture exam name, part, and year
#     exam_info_pattern = re.compile(r'(.+?)\s+(\(.*?\))\s+(\d{4})')
#
#     for para in doc.paragraphs:
#         text = para.text.strip()
#         if not text:
#             continue
#
#         if question_match := question_pattern.match(text):
#             question_number, initial_text, _ = question_match.groups()  # Exam year is parsed separately below
#             question_dict = {
#                 "question_number": question_number,
#                 "question_text": initial_text.strip(),
#                 "list_I": {"title": "", "options": {}},
#                 "list_II": {"title": "", "options": {}},
#                 "codes": [],
#                 "options": {},
#                 "answer": "",
#                 "exam_name": "",
#                 "exam_part": "",
#                 "exam_year": ""  # Added as a separate key
#             }
#
#             if list_match := list_title_pattern.search(text):
#                 list_I_title, list_II_title = list_match.groups()
#                 question_dict["list_I"]["title"] = list_I_title.strip()
#                 question_dict["list_II"]["title"] = list_II_title.strip()
#
#                 for match in list_I_option_pattern.finditer(text):
#                     letter, option = match.groups()
#                     question_dict["list_I"]["options"][letter] = option.strip()
#
#                 for match in list_II_option_pattern.finditer(text[list_match.end():]):
#                     number, option = match.groups()
#                     option = re.sub(r'\n.*', '', option)
#                     question_dict["list_II"]["options"][number] = option.strip()
#
#             if codes_match := codes_pattern.search(text):
#                 question_dict["codes"] = codes_match.group(1).split()
#
#             for match in options_pattern.finditer(text):
#                 code, option = match.groups()
#                 option = re.sub(r'\n.*', '', option)
#                 question_dict["options"][code] = option.strip()
#
#             if answer_match := answer_pattern.search(text):
#                 question_dict["answer"] = answer_match.group(1)
#
#             # Use exam_info_pattern to extract exam_name, exam_part, and exam_year
#             if exam_info_match := exam_info_pattern.search(text):
#                 question_dict["exam_name"] = exam_info_match.group(1).strip()
#                 question_dict["exam_part"] = exam_info_match.group(2).strip()
#                 question_dict["exam_year"] = exam_info_match.group(3).strip()  # Correctly extracting exam year
#
#             questions_list.append(question_dict)
#
#     return questions_list
#
# # Assuming 'sample_questions.docx' is the path to your Word document
# questions_json = parse_word_document('sample_questions.docx')
#
# # Convert the list of questions to a JSON string and print it
# print(json.dumps(questions_json, indent=4, ensure_ascii=False))
#
#
#
#
#
# def create_excel(data, excel_path):
#     wb = Workbook()
#     ws = wb.active
#
#     # Headers with placeholders for multiple options
#     headers = [
#         'Question Number', 'Question Text', 'Exam Year',
#         'List I Title', 'List I Option A', 'List I Option B', 'List I Option C', 'List I Option D',
#         'List II Title', 'List II Option 1', 'List II Option 2', 'List II Option 3', 'List II Option 4',
#         'Codes', 'Option A', 'Option B', 'Option C', 'Option D', 'Correct Answer',
#         'Exam Name', 'Exam Part'
#     ]
#     ws.append(headers)
#
#     for item in data:
#         list_I_options = item.get('list_I', {}).get('options', {})
#         list_II_options = item.get('list_II', {}).get('options', {})
#
#         # Prepare a row with placeholders for each option
#         row = [
#             item.get('question_number', ''),
#             item.get('question_text', ''),
#             item.get('exam_year', ''),
#             item.get('list_I', {}).get('title', ''),
#             list_I_options.get('A', ''), list_I_options.get('B', ''),
#             list_I_options.get('C', ''), list_I_options.get('D', ''),
#             item.get('list_II', {}).get('title', ''),
#             list_II_options.get('1', ''), list_II_options.get('2', ''),
#             list_II_options.get('3', ''), list_II_options.get('4', ''),
#             ', '.join(item.get('codes', [])),
#             item.get('options', {}).get('a', ''),
#             item.get('options', {}).get('b', ''),
#             item.get('options', {}).get('c', ''),
#             item.get('options', {}).get('d', ''),
#             item.get('answer', ''),
#             item.get('exam_name', ''),
#             item.get('exam_part', '')
#         ]
#         ws.append(row)
#
#     wb.save(excel_path)
#
# # Assuming the questions_json variable contains your data
# excel_file_path = 'parsed_questions_detailed02.xlsx'
# create_excel(questions_json, excel_file_path)
#
# print(excel_file_path)





import json
import re
from docx import Document
from openpyxl import Workbook

def parse_word_document(doc_path):
    doc = Document(doc_path)
    questions_list = []

    # Define regex patterns
    question_pattern = re.compile(r'^(\d+)\. (.+?)\[(\d{4})\]')
    list_title_pattern = re.compile(r'List-I\s+\((.*?)\)\s+List-II\s+\((.*?)\)')
    list_I_option_pattern = re.compile(r'([ABCD])\.\s+(.+?)(?=\s+[ABCD]\.\s+|\s+\d+\.\s+|\Z)', re.DOTALL)
    list_II_option_pattern = re.compile(r'(\d+)\.\s+(.+?)(?=\s+\d+\.\s+|\Z)', re.DOTALL)
    codes_pattern = re.compile(r'Codes:\s+([ABCD\s]+)')
    options_pattern = re.compile(r'\(([abcd])\)\s+(.*?)\s*(?=\([abcd]\)|$)', re.DOTALL)
    answer_pattern = re.compile(r'Answer\s+-\s+\((a|b|c|d)\)')
    exam_info_pattern = re.compile(r'(.+?)\s+(\(.*?\))\s+(\d{4})')

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if question_match := question_pattern.match(text):
            question_number, initial_text, _ = question_match.groups()
            question_dict = {
                "question_number": question_number,
                "question_text": initial_text.strip(),
                "list_I": {"title": "", "options": {}},
                "list_II": {"title": "", "options": {}},
                "codes": [],
                "options": {},
                "answer": "",
                "exam_name": "",
                "exam_part": "",
                "exam_year": "",
                "question_type": "mcqs_question"  # Default type
            }

            if list_match := list_title_pattern.search(text):
                # If List-I and List-II are found, it's a koat_question
                question_dict["question_type"] = "koat_question"

                list_I_title, list_II_title = list_match.groups()
                question_dict["list_I"]["title"] = list_I_title.strip()
                question_dict["list_II"]["title"] = list_II_title.strip()

                for match in list_I_option_pattern.finditer(text):
                    letter, option = match.groups()
                    question_dict["list_I"]["options"][letter] = option.strip()

                for match in list_II_option_pattern.finditer(text[list_match.end():]):
                    number, option = match.groups()
                    option = re.sub(r'\n.*', '', option)
                    question_dict["list_II"]["options"][number] = option.strip()

            if codes_match := codes_pattern.search(text):
                question_dict["codes"] = codes_match.group(1).split()

            for match in options_pattern.finditer(text):
                code, option = match.groups()
                option = re.sub(r'\n.*', '', option)
                question_dict["options"][code] = option.strip()

            if answer_match := answer_pattern.search(text):
                question_dict["answer"] = answer_match.group(1)

            if exam_info_match := exam_info_pattern.search(text):
                question_dict["exam_name"] = exam_info_match.group(1).strip()
                question_dict["exam_part"] = exam_info_match.group(2).strip()
                question_dict["exam_year"] = exam_info_match.group(3).strip()

            questions_list.append(question_dict)

    return questions_list

def create_excel(data, excel_path):
    wb = Workbook()
    ws = wb.active

    # Headers with placeholders for multiple options and added "Question Type"
    headers = [
        'Question Number', 'Question Text', 'Exam Year',
        'List I Title', 'List I Option A', 'List I Option B', 'List I Option C', 'List I Option D',
        'List II Title', 'List II Option 1', 'List II Option 2', 'List II Option 3', 'List II Option 4',
        'Codes', 'Option A', 'Option B', 'Option C', 'Option D', 'Correct Answer',
        'Exam Name', 'Exam Part', 'Question Type'
    ]
    ws.append(headers)

    for item in data:
        list_I_options = item.get('list_I', {}).get('options', {})
        list_II_options = item.get('list_II', {}).get('options', {})

        row = [
            item.get('question_number', ''),
            item.get('question_text', ''),
            item.get('exam_year', ''),
            item.get('list_I', {}).get('title', ''),
            list_I_options.get('A', ''), list_I_options.get('B', ''),
            list_I_options.get('C', ''), list_I_options.get('D', ''),
            item.get('list_II', {}).get('title', ''),
            list_II_options.get('1', ''), list_II_options.get('2', ''),
            list_II_options.get('3', ''), list_II_options.get('4', ''),
            ', '.join(item.get('codes', [])),
            item.get('options', {}).get('a', ''),
            item.get('options', {}).get('b', ''),
            item.get('options', {}).get('c', ''),
            item.get('options', {}).get('d', ''),
            item.get('answer', ''),
            item.get('exam_name', ''),
            item.get('exam_part', ''),
            item.get('question_type', '')  # Include the question type
        ]
        ws.append(row)

    wb.save(excel_path)

# Use the parse_word_document function to extract data from the DOCX file
questions_data = parse_word_document('sample_questions.docx')

# Use the create_excel function to create an Excel file from the parsed data
create_excel(questions_data, 'output_questions02.xlsx')

print('Excel file created successfully.')
