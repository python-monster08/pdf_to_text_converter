# from docx import Document
# import re

# # Define the path to the text file containing the extracted content
# file_path = 'OutputFiles/extracted_text2.txt'

# # Define a function to extract and format the content
# def extract_and_format(content):
#     # Regular expression patterns for each field
#     patterns = {
#         "[QUESTION_HEADING]": r"(\d+\.\s.*?\?)",
#         "[LIST_HEADING]": "",
#         "[LIST_HEADING_I]": "",
#         "[LIST_HEADING_II]": "",
#         "[LIST_CONTENT1]": "",
#         "[LIST_CONTENT1_A]": "",
#         "[LIST_CONTENT1_1]": "",
#         "[LIST_CONTENT2]": "",
#         "[LIST_CONTENT2_B]": "",
#         "[LIST_CONTENT2_2]": "",
#         "[LIST_CONTENT3]": "",
#         "[LIST_CONTENT3_C]": "",
#         "[LIST_CONTENT3_3]": "",
#         "[LIST_CONTENT4]": "",
#         "[LIST_CONTENT4_D]": "",
#         "[LIST_CONTENT4_4]": "",
#         "[CODE]": "",
#         "[OBJECTIVES]": r"\([a-d]\)\s[^\n]+",
#         "[OBJECTIVE1]": r"(\(a\).+?)\n",
#         "[OBJECTIVE2]": r"(\(b\).+?)\n",
#         "[OBJECTIVE3]": r"(\(c\).+?)\n",
#         "[OBJECTIVE4]": r"(\(d\).+?)\n",
#         "[EXAM_NAME]": "",
#         "[EXAM_TYPE]": "",
#         "[EXAM_YEAR]": r"\[(\d{4})\]",
#         "[CORRECT_ANSWER]": ""
#     }

#     # Process each question and its options
#     questions = re.findall(patterns["[QUESTION_HEADING]"], content, re.DOTALL)
#     options = {key: re.findall(pattern, content, re.DOTALL) for key, pattern in patterns.items() if key.startswith("[LIST_CONTENT")}

#     return questions, options

# # Load the content from the file
# with open(file_path, 'r') as file:
#     content = file.read()

# # Extract and format the content
# questions, options = extract_and_format(content)

# # Create a new Document
# doc = Document()

# # Add extracted content to the document
# for i, question in enumerate(questions):
#     doc.add_paragraph(f"[QUESTION_HEADING]: {question}")
#     for option_key in ["[LIST_CONTENT1]", "[LIST_CONTENT2]", "[LIST_CONTENT3]", "[LIST_CONTENT4]"]:
#         if i < len(options[option_key]):
#             doc.add_paragraph(f"{option_key}: {options[option_key][i]}")

#     # Add other fields to the document
#     doc.add_paragraph(f"[OBJECTIVES]: {options['[OBJECTIVES]'][i]}")
#     doc.add_paragraph(f"[EXAM_YEAR]: {options['[EXAM_YEAR]'][i]}")
#     doc.add_paragraph(f"[CORRECT_ANSWER]: {options['[CORRECT_ANSWER]'][i]}")
#     doc.add_paragraph("")  # Add an empty line between questions

# # Save the document
# output_docx_path = 'OutputFiles/Processed_Content.docx'
# doc.save(output_docx_path)

# print(f"Document saved at {output_docx_path}")



from docx import Document
import re

# Regular expressions for different types of MCQs
mcq_pattern = re.compile(r'(\d+\.\s.+?\?\s\[\d{4}\])(\n\([a-d]\)\s[^(\n]+)+', re.DOTALL)
mcq1_pattern = re.compile(r'(\d+\.\s.+?\[1995\])(\n\([a-d]\)\s.+?)+(\nSelect the correct answer using the codes given below:(?:\n\([a-d]\).+?)+)', re.DOTALL)

# Define the path to the text file containing the extracted content
file_path = 'OutputFiles/extracted_text2.txt'

# Function to process matches and save to a Word document
def process_matches(matches, file_name):
    doc = Document()
    for match in matches:
        # Add the question heading
        doc.add_paragraph(f"[QUESTION_HEADING]: {match[0]}")
        # Add options or additional information
        for part in match[1:]:
            doc.add_paragraph(part.strip())
        doc.add_paragraph("")  # Add an empty line between questions
    # Save the document
    doc.save(f'OutputFiles/{file_name}.docx')
    print(f"Document saved at OutputFiles/{file_name}.docx")

# Load the content from the file
with open(file_path, 'r') as file:
    content = file.read()

# Find matches for the first MCQ pattern and process them
mcq_matches = mcq_pattern.findall(content)
process_matches(mcq_matches, "MCQs_Matched_With_Pattern1")

# Find matches for the second MCQ pattern and process them
mcq1_matches = mcq1_pattern.findall(content)
process_matches(mcq1_matches, "MCQs_Matched_With_Pattern2")
