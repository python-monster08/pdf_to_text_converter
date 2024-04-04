import docx
import pdfplumber
def extract_questions_from_doc(file_name):
    # Create an instance of the Word document
    document = docx.Document(file_name)

    questions = []

    for paragraph in document.paragraphs:
        # Check if the paragraph has the question format
        if ':' in paragraph.text:
            question_text = paragraph.text.strip()

            question_data = {
                'question_type': 'mcqs',
            }

            objective_list = []
            correct_answer = None

            # Process the next paragraphs for the question
            for next_paragraph in document.paragraphs[document.paragraphs.index(paragraph) + 1:]:
                if any(bullet in next_paragraph.text for bullet in ['(a)', '(b)', '(c)', '(d)']):
                    objectives = next_paragraph.text.strip().split('\n')
                    objective_list = [option.strip() for option in objectives if option.strip()]

                    if 'correct answer' in next_paragraph.text.lower():
                        correct_answer_text = [option for option in objectives if 'correct answer' in option.lower()][0]
                        correct_answer = correct_answer_text.strip()[correct_answer_text.strip().index('(') + 1:correct_answer_text.strip().index(')')]
                        objective_list.remove(correct_answer_text)

                # If it's the next question, break the loop
                if not any(bullet in next_paragraph.text for bullet in ['(a)', '(b)', '(c)', '(d)']):
                    break

            question_data['question_text'] = question_text
            question_data['all_objectives'] = objective_list
            question_data['objective_a'] = objective_list[0]
            question_data['objective_b'] = objective_list[1]
            question_data['objective_c'] = objective_list[2]
            question_data['objective_d'] = objective_list[3]
            question_data['correct_answer'] = correct_answer

            questions.append(question_data)

    return questions

def extract_text_from_pdf(file_name):
    with pdfplumber.open(file_name) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
    return text

pdf_text = extract_text_from_pdf('History Medieval Ques Only Pages 15-16.pdf')
questions = extract_questions_from_doc(pdf_text)
print(questions)
# Example usage
questions = extract_questions_from_doc('History Medieval Ques Only Pages 15-16.pdf')
print(questions)